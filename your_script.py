import sys
import re
import os
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook, load_workbook
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QFileDialog, QListWidget, QListWidgetItem, QLabel,
    QComboBox, QSpinBox, QDoubleSpinBox, QCheckBox, QTabWidget,
    QMessageBox, QDialog, QDialogButtonBox, QGroupBox, QFormLayout,
    QMenu, QColorDialog, QFontComboBox
)
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QColor

# ================== Word COM 检测与转换 ==================

def check_word_available():
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Quit()
        return True
    except:
        return False

def convert_doc_to_docx(doc_path):
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(doc_path)
        temp_dir = tempfile.gettempdir()
        base_name = os.path.splitext(os.path.basename(doc_path))[0]
        docx_path = os.path.join(temp_dir, f"{base_name}_converted.docx")
        doc.SaveAs2(docx_path, FileFormat=16)
        doc.Close()
        word.Quit()
        return docx_path
    except:
        return None

# ================== 编号格式检测与生成 ==================

MAX_LEVEL = 5

def clean_title_text(text):
    text = text.strip()
    if not text:
        return text
    end_punctuation = '：:。；;，,、—…'
    while text and text[-1] in end_punctuation:
        text = text[:-1].rstrip()
    return text

def detect_number_format(text):
    text = text.strip()

    # 中文序号
    m = re.match(r'^([一二三四五六七八九十百千]+)([、.．]?)\s{0,4}(.*)', text)
    if m and (m.group(2) or m.group(3)):
        return ('chinese', m.group(2) or '、', False, m.group(1), m.group(3))

    # 括号类型
    m = re.match(r'^[（(]([一二三四五六七八九十百千]+|\d+|[A-Za-z]+|[IVXLCivxlc]+)[）)]\s{0,4}(.*)', text)
    if m:
        inner = m.group(1)
        rest = m.group(2)
        if re.match(r'^[一二三四五六七八九十百千]+$', inner):
            return ('chinese_paren', '', False, inner, rest)
        elif re.match(r'^\d+$', inner):
            return ('paren_digit', '', False, inner, rest)
        elif re.match(r'^[A-Za-z]+$', inner):
            return ('alpha_paren', '', False, inner, rest)
        elif re.match(r'^[IVXLCivxlc]+$', inner):
            return ('roman_paren', '', False, inner, rest)

    # 右括号结尾
    m = re.match(r'^(\d+|[A-Za-z]+|[IVXLCivxlc]+)([)）])\s{0,4}(.*)', text)
    if m:
        num = m.group(1)
        sep = m.group(2)
        rest = m.group(3)
        if num.isdigit():
            return ('paren_digit', sep, False, num, rest)
        elif re.match(r'^[A-Za-z]+$', num):
            return ('alpha_paren', sep, False, num, rest)
        elif re.match(r'^[IVXLCivxlc]+$', num):
            return ('roman_paren', sep, False, num, rest)

    # 圈号
    m = re.match(r'^([①-⑩])\s{0,4}(.*)', text)
    if m:
        return ('circle', '', False, m.group(1), m.group(2))

    # 数字编号
    m = re.match(r'^(\d+(?:[.．]\d+)*)([、.．]?)\s{0,4}(.*)', text)
    if m:
        return ('digit', m.group(2), False, m.group(1), m.group(3))

    # 罗马数字
    m = re.match(r'^([IVXLCivxlc]+(?:[.．、](?:[IVXLCivxlc]+|\d+))*)([、.．]?)\s{0,4}(.*)', text)
    if m:
        return ('roman', m.group(2), False, m.group(1), m.group(3))

    # 英文字母
    m = re.match(r'^([A-Za-z](?:[.．、](?:[A-Za-z]|\d+))*)([、.．]?)\s{0,4}(.*)', text)
    if m:
        return ('alpha', m.group(2), False, m.group(1), m.group(3))

    return None

def get_heading_level_from_style(para):
    try:
        style_name = para.style.name.lower()
        m = re.search(r'(heading|标题)\s*(\d+)', style_name)
        if m:
            level = int(m.group(2))
            return min(level, MAX_LEVEL)
    except:
        pass
    return None

def is_toc_start(para):
    text = para.text.strip()
    if not text:
        return False
    if text == '目录' or text.startswith('目录') or '目录' in text[:10]:
        return True
    if text.lower() in ['contents', 'table of contents']:
        return True
    return False

def is_heading_paragraph(para):
    text = para.text.strip()
    if not text:
        return False, 0

    level = get_heading_level_from_style(para)
    if level is not None:
        return True, level

    detected = detect_number_format(text)
    if not detected:
        return False, 0

    num_type, sep, has_space, num_part, rest = detected
    title_text = clean_title_text(rest)

    if not title_text:
        if num_type in ('digit', 'roman', 'alpha'):
            parts = re.findall(r'[.．]', num_part)
            level = len(parts) + 1
            return True, min(level, MAX_LEVEL)
        else:
            return False, 0

    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', title_text))
    english_words = len(re.findall(r'[A-Za-z]+', title_text))
    if chinese_chars > 30 or english_words > 20:
        return False, 0

    if num_type == 'digit':
        level = len(num_part.split('.'))
    elif num_type == 'roman':
        level = len(re.findall(r'[.．、]', num_part)) + 1
    elif num_type == 'alpha':
        level = len(re.findall(r'[.．、]', num_part)) + 1
    elif num_type == 'chinese':
        level = 1
    elif num_type in ('chinese_paren', 'paren_digit', 'alpha_paren', 'roman_paren'):
        level = 2
    elif num_type == 'circle':
        level = 3
    else:
        level = 1

    return True, min(level, MAX_LEVEL)

def int_to_roman(num):
    values = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    symbols = ['M', 'CM', 'D', 'CD', 'C', 'XC', 'L', 'XL', 'X', 'IX', 'V', 'IV', 'I']
    result = ''
    for i, value in enumerate(values):
        while num >= value:
            result += symbols[i]
            num -= value
    return result

def int_to_alpha(num, upper=True):
    result = ''
    while num > 0:
        num -= 1
        char = chr(ord('A' if upper else 'a') + num % 26)
        result = char + result
        num //= 26
    return result

def to_chinese_number(num):
    digits = ['零','一','二','三','四','五','六','七','八','九']
    if num <= 10:
        return digits[num] if num < 10 else '十'
    elif num < 20:
        return '十' + (digits[num%10] if num%10 else '')
    elif num < 100:
        return digits[num//10] + '十' + (digits[num%10] if num%10 else '')
    return str(num)

def generate_number(num_type, level, counters, sep, has_space):
    if num_type == 'digit':
        nums = [str(counters[i]) for i in range(level)]
        num_str = '.'.join(nums)
        if sep and level == 1:
            num_str += sep
        if has_space:
            num_str += ' '
        return num_str
    elif num_type == 'roman':
        nums = [int_to_roman(counters[i]) for i in range(level)]
        num_str = sep.join(nums) if sep else '.'.join(nums)
        if has_space:
            num_str += ' '
        return num_str
    elif num_type == 'alpha':
        nums = [int_to_alpha(counters[i]) for i in range(level)]
        num_str = sep.join(nums) if sep else '.'.join(nums)
        if has_space:
            num_str += ' '
        return num_str
    elif num_type == 'chinese':
        return to_chinese_number(counters[0]) + '、'
    elif num_type == 'chinese_paren':
        return '（' + to_chinese_number(counters[1]) + '）'
    elif num_type == 'paren_digit':
        return str(counters[1]) + '）'
    elif num_type == 'circle':
        circles = '①②③④⑤⑥⑦⑧⑨⑩'
        return circles[min(counters[2]-1, 9)] if counters[2] <= 10 else str(counters[2])
    else:
        num_str = '.'.join(str(counters[i]) for i in range(level))
        if has_space:
            num_str += ' '
        return num_str

def set_heading_style(paragraph, level):
    try:
        style_name = f'Heading {min(level, 9)}'
        paragraph.style = paragraph.part.document.styles[style_name]
    except:
        pass

def add_toc(paragraph, levels=5):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f'TOC \\o "1-{levels}" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = "目录：请在此处右键选择更新域以生成目录内容。"
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(t)
    run._r.append(fldChar3)

def get_table_name(table, doc, index):
    return f"表格 {index+1}"

def is_empty_paragraph(para):
    return not para.text.strip()

def count_empty_paragraphs_before(doc, para_index):
    count = 0
    idx = para_index - 1
    while idx >= 0 and is_empty_paragraph(doc.paragraphs[idx]):
        count += 1
        idx -= 1
    return count

def count_empty_paragraphs_after(doc, para_index):
    count = 0
    idx = para_index + 1
    while idx < len(doc.paragraphs) and is_empty_paragraph(doc.paragraphs[idx]):
        count += 1
        idx += 1
    return count

# ================== 主窗口 ==================

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Word 智能标题与表格处理工具")
        self.setGeometry(100, 100, 1200, 900)
        self.current_doc_path = None
        self.doc = None
        self.headings = []
        self.init_ui()

    def init_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)

        file_layout = QHBoxLayout()
        self.btn_open = QPushButton("打开 Word 文档")
        self.btn_open.clicked.connect(self.open_document)
        self.lbl_file = QLabel("未选择文件")
        file_layout.addWidget(self.btn_open)
        file_layout.addWidget(self.lbl_file, 1)
        main_layout.addLayout(file_layout)

        self.tabs = QTabWidget()
        main_layout.addWidget(self.tabs)

        self.tab1 = QWidget()
        self.tabs.addTab(self.tab1, "标题识别与目录")
        self.init_tab1()

        self.tab2 = QWidget()
        self.tabs.addTab(self.tab2, "表格填充")
        self.init_tab2()

        self.tab3 = QWidget()
        self.tabs.addTab(self.tab3, "Word ↔ Excel")
        self.init_tab3()

        self.statusBar().showMessage("就绪")

    def init_tab1(self):
        layout = QVBoxLayout(self.tab1)

        self.btn_scan = QPushButton("扫描标题")
        self.btn_scan.clicked.connect(self.scan_headings)
        layout.addWidget(self.btn_scan)

        skip_row = QHBoxLayout()
        skip_row.addWidget(QLabel("跳过前几页："))
        self.spin_skip_pages = QSpinBox()
        self.spin_skip_pages.setRange(0, 20)
        self.spin_skip_pages.setValue(0)
        skip_row.addWidget(self.spin_skip_pages)
        skip_row.addWidget(QLabel("页（0表示不跳过）"))
        layout.addLayout(skip_row)

        self.list_headings = QListWidget()
        self.list_headings.setSelectionMode(QListWidget.MultiSelection)
        self.list_headings.setContextMenuPolicy(Qt.CustomContextMenu)
        self.list_headings.customContextMenuRequested.connect(self.show_heading_context_menu)
        layout.addWidget(QLabel("识别到的标题（带[目录]前缀的为目录区域标题，可手动删除）："))
        layout.addWidget(self.list_headings)

        btn_row = QHBoxLayout()
        self.btn_add_heading = QPushButton("从段落中选取添加标题")
        self.btn_add_heading.clicked.connect(self.show_paragraph_picker)
        self.btn_delete_heading = QPushButton("删除选中标题")
        self.btn_delete_heading.clicked.connect(self.delete_selected_headings)
        self.btn_mark_non_heading = QPushButton("标记为非标题")
        self.btn_mark_non_heading.clicked.connect(self.mark_selected_as_non_heading)
        btn_row.addWidget(self.btn_add_heading)
        btn_row.addWidget(self.btn_delete_heading)
        btn_row.addWidget(self.btn_mark_non_heading)
        layout.addLayout(btn_row)

        level_row = QHBoxLayout()
        level_row.addWidget(QLabel("调整选中标题层级："))
        self.combo_set_level = QComboBox()
        self.combo_set_level.addItems([f"{i}级" for i in range(1, MAX_LEVEL+1)])
        level_row.addWidget(self.combo_set_level)
        self.btn_set_level = QPushButton("设为指定层级")
        self.btn_set_level.clicked.connect(self.set_selected_level)
        self.btn_level_up = QPushButton("提高一级")
        self.btn_level_up.clicked.connect(lambda: self.change_selected_level(-1))
        self.btn_level_down = QPushButton("降低一级")
        self.btn_level_down.clicked.connect(lambda: self.change_selected_level(1))
        level_row.addWidget(self.btn_level_up)
        level_row.addWidget(self.btn_level_down)
        layout.addLayout(level_row)

        self.btn_apply_style = QPushButton("将选中标题应用为 Word 标题样式")
        self.btn_apply_style.clicked.connect(self.apply_heading_styles)
        layout.addWidget(self.btn_apply_style)

        group_format = QGroupBox("标题格式统一与自动编号")
        format_layout = QVBoxLayout()

        level_select_row = QHBoxLayout()
        level_select_row.addWidget(QLabel("作用层级："))
        self.check_unify_levels = []
        for i in range(MAX_LEVEL):
            check = QCheckBox(f"{i+1}级")
            check.setChecked(True)
            self.check_unify_levels.append(check)
            level_select_row.addWidget(check)
        format_layout.addLayout(level_select_row)

        num_row = QHBoxLayout()
        num_row.addWidget(QLabel("编号样式："))
        self.combo_number_style = QComboBox()
        self.combo_number_style.addItems([
            "与完美格式标题一致",
            "数字编号：1, 1.1, 1.1.1",
            "罗马数字：I, I.I, I.I.I",
            "英文字母：A, A.a, A.a.a",
            "中文序号：一、（一）、1）、①",
            "保持原编号"
        ])
        num_row.addWidget(self.combo_number_style)
        format_layout.addLayout(num_row)

        source_row = QHBoxLayout()
        source_row.addWidget(QLabel("格式来源："))
        self.radio_format_source = QComboBox()
        self.radio_format_source.addItems(["从选定标题复制格式", "手动设置格式"])
        source_row.addWidget(self.radio_format_source)
        format_layout.addLayout(source_row)

        self.manual_widget = QWidget()
        manual_layout = QVBoxLayout()
        row1 = QHBoxLayout()
        row1.addWidget(QLabel("字体："))
        self.combo_font = QFontComboBox()
        row1.addWidget(self.combo_font)
        row1.addWidget(QLabel("字号："))
        self.combo_font_size = QComboBox()
        self.combo_font_size.addItems(["9","10","10.5","11","12","14","16","18","20","22","24","26","28","36"])
        self.combo_font_size.setCurrentText("16")
        row1.addWidget(self.combo_font_size)
        manual_layout.addLayout(row1)

        row2 = QHBoxLayout()
        self.check_bold = QCheckBox("加粗")
        self.check_bold.setChecked(True)
        self.check_italic = QCheckBox("斜体")
        self.check_color = QCheckBox("字体颜色")
        self.btn_color = QPushButton("选择颜色")
        self.btn_color.clicked.connect(self.pick_color)
        self.color_value = None
        row2.addWidget(self.check_bold)
        row2.addWidget(self.check_italic)
        row2.addWidget(self.check_color)
        row2.addWidget(self.btn_color)
        manual_layout.addLayout(row2)

        row3 = QHBoxLayout()
        row3.addWidget(QLabel("对齐："))
        self.combo_alignment = QComboBox()
        self.combo_alignment.addItems(["左对齐", "居中", "右对齐", "两端对齐"])
        row3.addWidget(self.combo_alignment)
        manual_layout.addLayout(row3)

        row4 = QHBoxLayout()
        row4.addWidget(QLabel("段前(磅)："))
        self.spin_space_before = QDoubleSpinBox()
        self.spin_space_before.setRange(0, 100)
        row4.addWidget(self.spin_space_before)
        row4.addWidget(QLabel("段后(磅)："))
        self.spin_space_after = QDoubleSpinBox()
        self.spin_space_after.setRange(0, 100)
        row4.addWidget(self.spin_space_after)
        manual_layout.addLayout(row4)

        row5 = QHBoxLayout()
        row5.addWidget(QLabel("行距："))
        self.combo_line_spacing = QComboBox()
        self.combo_line_spacing.addItems(["单倍行距", "1.5倍行距", "2倍行距"])
        row5.addWidget(self.combo_line_spacing)
        row5.addWidget(QLabel("首行缩进(磅)："))
        self.spin_first_indent = QDoubleSpinBox()
        self.spin_first_indent.setRange(0, 100)
        row5.addWidget(self.spin_first_indent)
        manual_layout.addLayout(row5)

        row6 = QHBoxLayout()
        row6.addWidget(QLabel("左缩进(磅)："))
        self.spin_left_indent = QDoubleSpinBox()
        self.spin_left_indent.setRange(0, 100)
        row6.addWidget(self.spin_left_indent)
        row6.addWidget(QLabel("前空行数："))
        self.spin_empty_before = QSpinBox()
        self.spin_empty_before.setRange(0, 10)
        row6.addWidget(self.spin_empty_before)
        row6.addWidget(QLabel("后空行数："))
        self.spin_empty_after = QSpinBox()
        self.spin_empty_after.setRange(0, 10)
        row6.addWidget(self.spin_empty_after)
        manual_layout.addLayout(row6)

        self.manual_widget.setLayout(manual_layout)
        self.manual_widget.setVisible(False)
        format_layout.addWidget(self.manual_widget)

        self.radio_format_source.currentIndexChanged.connect(
            lambda idx: self.manual_widget.setVisible(idx == 1)
        )

        self.btn_unify = QPushButton("统一格式并自动编号")
        self.btn_unify.clicked.connect(self.unify_format_and_number)
        format_layout.addWidget(self.btn_unify)

        group_format.setLayout(format_layout)
        layout.addWidget(group_format)

        group_dir = QGroupBox("生成目录")
        dir_layout = QFormLayout()
        self.spin_toc_levels = QSpinBox()
        self.spin_toc_levels.setRange(1, MAX_LEVEL)
        self.spin_toc_levels.setValue(MAX_LEVEL)
        dir_layout.addRow("显示级别：", self.spin_toc_levels)
        self.btn_insert_toc = QPushButton("插入目录（含目录标题）")
        self.btn_insert_toc.clicked.connect(self.insert_toc)
        dir_layout.addRow(self.btn_insert_toc)
        group_dir.setLayout(dir_layout)
        layout.addWidget(group_dir)

        self.btn_generate = QPushButton("生成新 Word 文件")
        self.btn_generate.clicked.connect(self.generate_new_document)
        self.btn_generate.setStyleSheet("QPushButton { background-color: #4CAF50; color: white; font-size: 14px; padding: 8px; }")
        layout.addWidget(self.btn_generate)

    def init_tab2(self):
        layout = QVBoxLayout(self.tab2)
        self.btn_list_tables = QPushButton("列出所有表格")
        self.btn_list_tables.clicked.connect(self.list_tables)
        layout.addWidget(self.btn_list_tables)
        self.combo_tables = QComboBox()
        self.combo_tables.currentIndexChanged.connect(self.on_table_selected)
        layout.addWidget(QLabel("选择要填充的表格："))
        layout.addWidget(self.combo_tables)
        self.combo_seq_col = QComboBox()
        self.combo_title_col = QComboBox()
        layout.addWidget(QLabel("序号列："))
        layout.addWidget(self.combo_seq_col)
        layout.addWidget(QLabel("标题列："))
        layout.addWidget(self.combo_title_col)

        level_row = QHBoxLayout()
        level_row.addWidget(QLabel("填充层级："))
        self.check_fill_levels = []
        for i in range(MAX_LEVEL):
            check = QCheckBox(f"{i+1}级")
            check.setChecked(True)
            self.check_fill_levels.append(check)
            level_row.addWidget(check)
        layout.addLayout(level_row)

        self.btn_fill_table = QPushButton("填充表格")
        self.btn_fill_table.clicked.connect(self.fill_table)
        layout.addWidget(self.btn_fill_table)

    def init_tab3(self):
        layout = QVBoxLayout(self.tab3)
        group_export = QGroupBox("Word 表格导出到 Excel")
        export_layout = QVBoxLayout()
        self.btn_export = QPushButton("选择表格并导出")
        self.btn_export.clicked.connect(self.export_to_excel)
        export_layout.addWidget(self.btn_export)
        group_export.setLayout(export_layout)
        layout.addWidget(group_export)

        group_import = QGroupBox("Excel 数据导入 Word 表格")
        import_layout = QVBoxLayout()
        self.btn_import = QPushButton("选择 Excel 文件并导入")
        self.btn_import.clicked.connect(self.import_from_excel)
        import_layout.addWidget(self.btn_import)
        group_import.setLayout(import_layout)
        layout.addWidget(group_import)

    # ================== 文件操作 ==================
    def open_document(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择 Word 文档", "", "Word 文档 (*.docx *.doc)"
        )
        if not file_path:
            return

        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.doc':
            if check_word_available():
                QMessageBox.information(self, "提示", "检测到 .doc 文件，正在使用 Word 转换...")
                docx_path = convert_doc_to_docx(file_path)
                if docx_path and os.path.exists(docx_path):
                    self.current_doc_path = file_path
                    self.doc = Document(docx_path)
                    self.lbl_file.setText(f"{file_path} (已自动转换为 .docx)")
                    self.statusBar().showMessage("已加载（通过 Word 转换）")
                else:
                    QMessageBox.warning(self, "警告", "转换失败，请手动用 WPS 另存为 .docx 后重新打开")
                    return
            else:
                QMessageBox.warning(
                    self, "提示",
                    "未检测到 Microsoft Word。\n\n请用 WPS 打开该 .doc 文件，"
                    "然后另存为 .docx 格式，再重新打开 .docx 文件。"
                )
                return
        else:
            self.current_doc_path = file_path
            self.doc = Document(file_path)
            self.lbl_file.setText(file_path)
            self.statusBar().showMessage(f"已打开：{file_path}")

        self.headings = []
        self.list_headings.clear()
        self.combo_tables.clear()

    def get_new_save_path(self):
        if not self.current_doc_path:
            QMessageBox.warning(self, "警告", "请先打开 Word 文档")
            return None
        dir_path = os.path.dirname(self.current_doc_path)
        base_name = os.path.splitext(os.path.basename(self.current_doc_path))[0]
        new_name = f"{base_name}_已处理.docx"
        new_path = os.path.join(dir_path, new_name)
        if os.path.exists(new_path):
            new_path, _ = QFileDialog.getSaveFileName(self, "保存新文件", new_path, "Word 文档 (*.docx)")
            if not new_path:
                return None
        return new_path

    def generate_new_document(self):
        if not self.doc:
            QMessageBox.warning(self, "警告", "请先打开 Word 文档")
            return
        reply = QMessageBox.question(
            self, "确认生成",
            "确认所有设置已完成，生成新的 Word 文件？\n\n原文件不会被修改。",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply != QMessageBox.Yes:
            return
        new_path = self.get_new_save_path()
        if new_path:
            self.doc.save(new_path)
            QMessageBox.information(self, "完成", f"新文件已生成：{new_path}")
            self.current_doc_path = new_path
            self.doc = Document(new_path)
            self.scan_headings()

    # ================== 标题扫描与管理 ==================
    def scan_headings(self):
        if not self.doc:
            QMessageBox.warning(self, "警告", "请先打开 Word 文档")
            return
        self.headings = []
        self.list_headings.clear()

        skip_pages = self.spin_skip_pages.value()
        skip_paragraphs = skip_pages * 40

        in_toc = False
        toc_started = False

        for i, para in enumerate(self.doc.paragraphs):
            if i < skip_paragraphs:
                continue

            if not toc_started and is_toc_start(para):
                in_toc = True
                toc_started = True
                continue

            is_heading, level = is_heading_paragraph(para)
            if is_heading:
                detected = detect_number_format(para.text.strip())
                if detected:
                    num_type, sep, has_space, num_part, rest = detected
                    cleaned_text = clean_title_text(rest)
                    display_text = f"{num_part}{sep} {cleaned_text}".strip() if cleaned_text else num_part
                else:
                    display_text = para.text.strip()

                prefix = "[目录] " if in_toc else ""
                self.headings.append((i, level, display_text))
                item = QListWidgetItem(f"{prefix}[{level}] {display_text}")
                item.setData(Qt.UserRole, len(self.headings)-1)
                self.list_headings.addItem(item)

        if not self.headings:
            QMessageBox.information(self, "提示", "未识别到标题，请手动从段落中选取添加标题。")
        else:
            self.statusBar().showMessage(f"识别到 {len(self.headings)} 个标题，目录区域标题已标记为[目录]")

    def show_heading_context_menu(self, pos):
        menu = QMenu()
        act_add = menu.addAction("从段落中选取添加标题")
        act_delete = menu.addAction("删除选中标题")
        act_mark = menu.addAction("标记为非标题")
        menu.addSeparator()
        for i in range(1, MAX_LEVEL+1):
            action = menu.addAction(f"设为{i}级标题")
            action.triggered.connect(lambda checked, level=i: self.set_selected_level_value(level))
        menu.addSeparator()
        act_up = menu.addAction("提高一级")
        act_down = menu.addAction("降低一级")
        action = menu.exec_(self.list_headings.mapToGlobal(pos))
        if action == act_add:
            self.show_paragraph_picker()
        elif action == act_delete:
            self.delete_selected_headings()
        elif action == act_mark:
            self.mark_selected_as_non_heading()
        elif action == act_up:
            self.change_selected_level(-1)
        elif action == act_down:
            self.change_selected_level(1)

    def show_paragraph_picker(self):
        if not self.doc:
            QMessageBox.warning(self, "警告", "请先打开 Word 文档")
            return
        dialog = QDialog(self)
        dialog.setWindowTitle("选择要标记为标题的段落")
        dialog.setGeometry(200, 200, 600, 500)
        layout = QVBoxLayout(dialog)
        layout.addWidget(QLabel("勾选要标记为标题的段落："))
        list_widget = QListWidget()
        list_widget.setSelectionMode(QListWidget.MultiSelection)
        for i, para in enumerate(self.doc.paragraphs):
            if not para.text.strip():
                continue
            item = QListWidgetItem(f"{i+1}: {para.text.strip()[:60]}")
            item.setData(Qt.UserRole, i)
            list_widget.addItem(item)
        layout.addWidget(list_widget)
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)
        if dialog.exec_() != QDialog.Accepted:
            return
        selected = [item.data(Qt.UserRole) for item in list_widget.selectedItems()]
        if not selected:
            return
        level_dialog = QDialog(self)
        level_dialog.setWindowTitle("选择标题层级")
        level_layout = QFormLayout(level_dialog)
        combo_level = QComboBox()
        combo_level.addItems([f"{i}级" for i in range(1, MAX_LEVEL+1)])
        level_layout.addRow("层级：", combo_level)
        level_buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        level_buttons.accepted.connect(level_dialog.accept)
        level_buttons.rejected.connect(level_dialog.reject)
        level_layout.addRow(level_buttons)
        if level_dialog.exec_() != QDialog.Accepted:
            return
        level = combo_level.currentIndex() + 1
        for idx in selected:
            para = self.doc.paragraphs[idx]
            self.headings.append((idx, level, para.text.strip()))
            item = QListWidgetItem(f"[{level}] {para.text.strip()}")
            item.setData(Qt.UserRole, len(self.headings)-1)
            self.list_headings.addItem(item)
        self.statusBar().showMessage(f"已添加 {len(selected)} 个标题")

    def delete_selected_headings(self):
        selected_items = self.list_headings.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "警告", "请先选择要删除的标题")
            return
        reply = QMessageBox.question(self, "确认", "是否同时清除文档中这些段落的标题格式？",
                                     QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel)
        if reply == QMessageBox.Cancel:
            return
        clear_format = (reply == QMessageBox.Yes)
        indices_to_remove = []
        for item in selected_items:
            idx = item.data(Qt.UserRole)
            indices_to_remove.append(idx)
            if clear_format:
                para_idx = self.headings[idx][0]
                self.doc.paragraphs[para_idx].style = self.doc.styles['Normal']
        indices_to_remove.sort(reverse=True)
        for idx in indices_to_remove:
            self.headings.pop(idx)
        self.refresh_heading_list()

    def mark_selected_as_non_heading(self):
        selected_items = self.list_headings.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "警告", "请先选择要标记的标题")
            return
        reply = QMessageBox.question(self, "确认", "是否同时清除文档中这些段落的标题格式？",
                                     QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel)
        if reply == QMessageBox.Cancel:
            return
        clear_format = (reply == QMessageBox.Yes)
        indices_to_remove = []
        for item in selected_items:
            idx = item.data(Qt.UserRole)
            indices_to_remove.append(idx)
            if clear_format:
                para_idx = self.headings[idx][0]
                self.doc.paragraphs[para_idx].style = self.doc.styles['Normal']
        indices_to_remove.sort(reverse=True)
        for idx in indices_to_remove:
            self.headings.pop(idx)
        self.refresh_heading_list()

    def set_selected_level(self):
        level = self.combo_set_level.currentIndex() + 1
        self.set_selected_level_value(level)

    def set_selected_level_value(self, level):
        selected_items = self.list_headings.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "警告", "请先选择标题")
            return
        for item in selected_items:
            idx = item.data(Qt.UserRole)
            para_idx, _, text = self.headings[idx]
            self.headings[idx] = (para_idx, level, text)
        self.refresh_heading_list()

    def change_selected_level(self, delta):
        selected_items = self.list_headings.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "警告", "请先选择标题")
            return
        for item in selected_items:
            idx = item.data(Qt.UserRole)
            para_idx, level, text = self.headings[idx]
            new_level = max(1, min(MAX_LEVEL, level + delta))
            self.headings[idx] = (para_idx, new_level, text)
        self.refresh_heading_list()

    def refresh_heading_list(self):
        self.list_headings.clear()
        for i, (para_idx, level, text) in enumerate(self.headings):
            item = QListWidgetItem(f"[{level}] {text}")
            item.setData(Qt.UserRole, i)
            self.list_headings.addItem(item)

    def apply_heading_styles(self):
        selected_items = self.list_headings.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "警告", "请先选择要应用样式的标题")
            return
        for item in selected_items:
            idx = item.data(Qt.UserRole)
            para_idx, level, _ = self.headings[idx]
            set_heading_style(self.doc.paragraphs[para_idx], level)
        self.statusBar().showMessage("已应用标题样式，点击生成新 Word 文件保存")

    # ================== 格式统一与自动编号 ==================
    def pick_color(self):
        color = QColorDialog.getColor()
        if color.isValid():
            self.color_value = color
            self.btn_color.setStyleSheet(f"background-color: {color.name()}")
            self.check_color.setChecked(True)

    def get_selected_levels(self):
        levels = []
        for i, check in enumerate(self.check_unify_levels):
            if check.isChecked():
                levels.append(i+1)
        return levels

    def get_paragraph_full_format(self, para, para_index):
        format_info = {}
        if para.runs:
            run = para.runs[0]
            format_info['font_name'] = run.font.name
            format_info['font_size'] = run.font.size.pt if run.font.size else None
            format_info['bold'] = run.font.bold
            format_info['italic'] = run.font.italic
            format_info['color'] = run.font.color.rgb if run.font.color and run.font.color.rgb else None
        else:
            style = para.style
            format_info['font_name'] = style.font.name
            format_info['font_size'] = style.font.size.pt if style.font.size else None
            format_info['bold'] = style.font.bold
            format_info['italic'] = style.font.italic
            format_info['color'] = style.font.color.rgb if style.font.color and style.font.color.rgb else None

        pf = para.paragraph_format
        format_info['alignment'] = pf.alignment
        format_info['space_before'] = pf.space_before.pt if pf.space_before else 0
        format_info['space_after'] = pf.space_after.pt if pf.space_after else 0
        format_info['line_spacing'] = pf.line_spacing
        format_info['line_spacing_rule'] = pf.line_spacing_rule
        format_info['first_line_indent'] = pf.first_line_indent.pt if pf.first_line_indent else 0
        format_info['left_indent'] = pf.left_indent.pt if pf.left_indent else 0
        format_info['right_indent'] = pf.right_indent.pt if pf.right_indent else 0
        format_info['keep_with_next'] = pf.keep_with_next
        format_info['keep_together'] = pf.keep_together
        format_info['page_break_before'] = pf.page_break_before
        format_info['widow_control'] = pf.widow_control
        format_info['empty_before'] = count_empty_paragraphs_before(self.doc, para_index)
        format_info['empty_after'] = count_empty_paragraphs_after(self.doc, para_index)
        return format_info

    def apply_paragraph_full_format(self, para, para_index, format_info):
        for run in para.runs:
            if format_info.get('font_name'):
                run.font.name = format_info['font_name']
                if run._element.rPr is not None:
                    rFonts = run._element.rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        run._element.rPr.append(rFonts)
                    rFonts.set(qn('w:eastAsia'), format_info['font_name'])
            if format_info.get('font_size'):
                run.font.size = Pt(format_info['font_size'])
            if format_info.get('bold') is not None:
                run.font.bold = format_info['bold']
            if format_info.get('italic') is not None:
                run.font.italic = format_info['italic']
            if format_info.get('color'):
                run.font.color.rgb = format_info['color']

        pf = para.paragraph_format
        if format_info.get('alignment') is not None:
            pf.alignment = format_info['alignment']
        pf.space_before = Pt(format_info.get('space_before', 0))
        pf.space_after = Pt(format_info.get('space_after', 0))
        if format_info.get('line_spacing') is not None:
            pf.line_spacing = format_info['line_spacing']
        if format_info.get('line_spacing_rule') is not None:
            pf.line_spacing_rule = format_info['line_spacing_rule']
        pf.first_line_indent = Pt(format_info.get('first_line_indent', 0))
        pf.left_indent = Pt(format_info.get('left_indent', 0))
        if format_info.get('right_indent') is not None:
            pf.right_indent = Pt(format_info['right_indent'])
        if format_info.get('keep_with_next') is not None:
            pf.keep_with_next = format_info['keep_with_next']
        if format_info.get('keep_together') is not None:
            pf.keep_together = format_info['keep_together']
        if format_info.get('page_break_before') is not None:
            pf.page_break_before = format_info['page_break_before']
        if format_info.get('widow_control') is not None:
            pf.widow_control = format_info['widow_control']

        target_before = format_info.get('empty_before', 0)
        target_after = format_info.get('empty_after', 0)
        self.adjust_empty_paragraphs(para_index, target_before, target_after)

    def adjust_empty_paragraphs(self, para_index, target_before, target_after):
        current_before = count_empty_paragraphs_before(self.doc, para_index)
        if current_before > target_before:
            for _ in range(current_before - target_before):
                idx = para_index - 1
                if idx >= 0 and is_empty_paragraph(self.doc.paragraphs[idx]):
                    self.doc.paragraphs[idx]._element.getparent().remove(self.doc.paragraphs[idx]._element)
                    para_index -= 1
        elif current_before < target_before:
            for _ in range(target_before - current_before):
                new_para = self.doc.paragraphs[para_index].insert_paragraph_before()
                new_para.style = self.doc.styles['Normal']

        current_after = count_empty_paragraphs_after(self.doc, para_index)
        if current_after > target_after:
            for _ in range(current_after - target_after):
                idx = para_index + 1
                if idx < len(self.doc.paragraphs) and is_empty_paragraph(self.doc.paragraphs[idx]):
                    self.doc.paragraphs[idx]._element.getparent().remove(self.doc.paragraphs[idx]._element)
        elif current_after < target_after:
            for _ in range(target_after - current_after):
                new_para = self.doc.paragraphs[para_index].insert_paragraph_before()
                new_para.style = self.doc.styles['Normal']
                new_para._element.addnext(self.doc.paragraphs[para_index]._element)

    def apply_manual_format_to_para(self, para, para_index):
        font_name = self.combo_font.currentFont().family()
        font_size = float(self.combo_font_size.currentText())
        for run in para.runs:
            run.font.name = font_name
            if run._element.rPr is not None:
                rFonts = run._element.rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    run._element.rPr.append(rFonts)
                rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)
            if self.check_bold.isChecked():
                run.font.bold = True
            if self.check_italic.isChecked():
                run.font.italic = True
            if self.check_color.isChecked() and self.color_value:
                run.font.color.rgb = RGBColor(self.color_value.red(), self.color_value.green(), self.color_value.blue())

        pf = para.paragraph_format
        align_map = {"左对齐": WD_ALIGN_PARAGRAPH.LEFT, "居中": WD_ALIGN_PARAGRAPH.CENTER,
                     "右对齐": WD_ALIGN_PARAGRAPH.RIGHT, "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY}
        pf.alignment = align_map.get(self.combo_alignment.currentText(), WD_ALIGN_PARAGRAPH.LEFT)
        pf.space_before = Pt(self.spin_space_before.value())
        pf.space_after = Pt(self.spin_space_after.value())
        ls_map = {"单倍行距": 1.0, "1.5倍行距": 1.5, "2倍行距": 2.0}
        ls_text = self.combo_line_spacing.currentText()
        if ls_text in ls_map:
            pf.line_spacing = ls_map[ls_text]
        pf.first_line_indent = Pt(self.spin_first_indent.value())
        pf.left_indent = Pt(self.spin_left_indent.value())
        self.adjust_empty_paragraphs(para_index, self.spin_empty_before.value(), self.spin_empty_after.value())

    def unify_format_and_number(self):
        if not self.doc or not self.headings:
            QMessageBox.warning(self, "警告", "请先打开文档并扫描标题")
            return
        levels_to_unify = self.get_selected_levels()
        if not levels_to_unify:
            QMessageBox.warning(self, "警告", "请选择至少一个层级")
            return
        format_source = self.radio_format_source.currentIndex()
        if format_source == 0:
            selected_items = self.list_headings.selectedItems()
            if len(selected_items) != 1:
                QMessageBox.warning(self, "警告", "请选择一个完美格式标题")
                return
            source_idx = selected_items[0].data(Qt.UserRole)
            source_para_idx, source_level, _ = self.headings[source_idx]
            source_para = self.doc.paragraphs[source_para_idx]
            source_format = self.get_paragraph_full_format(source_para, source_para_idx)
            for i, (para_idx, level, text) in enumerate(self.headings):
                if level in levels_to_unify:
                    self.apply_paragraph_full_format(self.doc.paragraphs[para_idx], para_idx, source_format)
        else:
            for i, (para_idx, level, text) in enumerate(self.headings):
                if level in levels_to_unify:
                    self.apply_manual_format_to_para(self.doc.paragraphs[para_idx], para_idx)

        number_style = self.combo_number_style.currentIndex()
        self.auto_number_headings(number_style)
        self.statusBar().showMessage("格式统一与自动编号完成，点击生成新 Word 文件保存")

    def auto_number_headings(self, number_style):
        if number_style == 0:
            self.number_by_perfect_format()
        elif number_style == 1:
            self.number_by_digit()
        elif number_style == 2:
            self.number_by_roman()
        elif number_style == 3:
            self.number_by_alpha()
        elif number_style == 4:
            self.number_by_chinese()

    def number_by_perfect_format(self):
        selected_items = self.list_headings.selectedItems()
        if not selected_items:
            return
        source_idx = selected_items[0].data(Qt.UserRole)
        source_para_idx, source_level, source_text = self.headings[source_idx]
        detected = detect_number_format(source_text)
        if not detected:
            return
        num_type, sep, has_space, num_part, rest = detected

        counters = [0] * MAX_LEVEL
        for i, (para_idx, level, text) in enumerate(self.headings):
            counters[level-1] += 1
            for j in range(level, MAX_LEVEL):
                counters[j] = 0
            new_num = generate_number(num_type, level, counters, sep, has_space)
            title_text = self.remove_old_number(text)
            new_text = f"{new_num}{title_text}"
            self.set_paragraph_text(self.doc.paragraphs[para_idx], new_text)
            self.headings[i] = (para_idx, level, new_text)

    def number_by_digit(self):
        counters = [0] * MAX_LEVEL
        for i, (para_idx, level, text) in enumerate(self.headings):
            counters[level-1] += 1
            for j in range(level, MAX_LEVEL):
                counters[j] = 0
            new_num = '.'.join(str(counters[j]) for j in range(level))
            new_text = f"{new_num} {self.remove_old_number(text)}"
            self.set_paragraph_text(self.doc.paragraphs[para_idx], new_text)
            self.headings[i] = (para_idx, level, new_text)

    def number_by_roman(self):
        counters = [0] * MAX_LEVEL
        for i, (para_idx, level, text) in enumerate(self.headings):
            counters[level-1] += 1
            for j in range(level, MAX_LEVEL):
                counters[j] = 0
            nums = [int_to_roman(counters[j]) for j in range(level)]
            new_num = '.'.join(nums)
            new_text = f"{new_num} {self.remove_old_number(text)}"
            self.set_paragraph_text(self.doc.paragraphs[para_idx], new_text)
            self.headings[i] = (para_idx, level, new_text)

    def number_by_alpha(self):
        counters = [0] * MAX_LEVEL
        for i, (para_idx, level, text) in enumerate(self.headings):
            counters[level-1] += 1
            for j in range(level, MAX_LEVEL):
                counters[j] = 0
            nums = [int_to_alpha(counters[j]) for j in range(level)]
            new_num = '.'.join(nums)
            new_text = f"{new_num} {self.remove_old_number(text)}"
            self.set_paragraph_text(self.doc.paragraphs[para_idx], new_text)
            self.headings[i] = (para_idx, level, new_text)

    def number_by_chinese(self):
        counters = [0] * MAX_LEVEL
        for i, (para_idx, level, text) in enumerate(self.headings):
            counters[level-1] += 1
            for j in range(level, MAX_LEVEL):
                counters[j] = 0
            if level == 1:
                new_num = to_chinese_number(counters[0]) + '、'
            elif level == 2:
                new_num = '（' + to_chinese_number(counters[1]) + '）'
            elif level == 3:
                new_num = str(counters[2]) + '）'
            elif level == 4:
                circles = '①②③④⑤⑥⑦⑧⑨⑩'
                new_num = circles[min(counters[3]-1, 9)]
            else:
                new_num = int_to_alpha(counters[4])
            new_text = f"{new_num}{self.remove_old_number(text)}"
            self.set_paragraph_text(self.doc.paragraphs[para_idx], new_text)
            self.headings[i] = (para_idx, level, new_text)

    def remove_old_number(self, text):
        detected = detect_number_format(text)
        if detected:
            num_type, sep, has_space, num_part, rest = detected
            return clean_title_text(rest)
        return clean_title_text(text)

    def set_paragraph_text(self, para, new_text):
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.add_run(new_text)

    # ================== 目录生成 ==================
    def insert_toc(self):
        if not self.doc:
            QMessageBox.warning(self, "警告", "请先打开 Word 文档")
            return
        levels = self.spin_toc_levels.value()
        first_para = self.doc.paragraphs[0]
        toc_para = first_para.insert_paragraph_before()
        add_toc(toc_para, levels)
        title_para = toc_para.insert_paragraph_before()
        title_run = title_para.add_run("目录")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.statusBar().showMessage("目录已插入，点击生成新 Word 文件保存")

    # ================== 表格填充 ==================
    def list_tables(self):
        if not self.doc:
            QMessageBox.warning(self, "警告", "请先打开 Word 文档")
            return
        self.combo_tables.clear()
        for i, table in enumerate(self.doc.tables):
            self.combo_tables.addItem(f"{i+1}: {get_table_name(table, self.doc, i)}", i)
        if not self.doc.tables:
            QMessageBox.information(self, "提示", "文档中没有表格")
        else:
            self.statusBar().showMessage(f"找到 {len(self.doc.tables)} 个表格")

    def on_table_selected(self, index):
        if index < 0 or not self.doc:
            return
        table_idx = self.combo_tables.itemData(index)
        if table_idx is None:
            return
        table = self.doc.tables[table_idx]
        self.combo_seq_col.clear()
        self.combo_title_col.clear()
        headers = []
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                headers.append(cell.text.strip())
        for col_idx in range(len(table.columns)):
            label = f"{col_idx+1} - {headers[col_idx]}" if col_idx < len(headers) and headers[col_idx] else f"列 {col_idx+1}"
            self.combo_seq_col.addItem(label, col_idx)
            self.combo_title_col.addItem(label, col_idx)

    def fill_table(self):
        if not self.doc or not self.headings:
            QMessageBox.warning(self, "警告", "请先打开文档并扫描标题")
            return
        table_idx = self.combo_tables.currentData()
        if table_idx is None:
            QMessageBox.warning(self, "警告", "请先选择表格")
            return
        seq_col = self.combo_seq_col.currentData()
        title_col = self.combo_title_col.currentData()
        if seq_col is None or title_col is None:
            QMessageBox.warning(self, "警告", "请指定列")
            return
        levels = []
        for i, check in enumerate(self.check_fill_levels):
            if check.isChecked():
                levels.append(i+1)
        titles = [(level, text) for _, level, text in self.headings if level in levels]
        if not titles:
            QMessageBox.warning(self, "警告", "没有符合层级的标题")
            return
        table = self.doc.tables[table_idx]
        needed = len(titles)
        current = len(table.rows) - 1
        if current < needed:
            for _ in range(needed - current):
                table.add_row()
        for i, (level, text) in enumerate(titles):
            row = table.rows[i+1]
            row.cells[seq_col].text = self.extract_seq_text(text, i+1)
            row.cells[title_col].text = text
        self.statusBar().showMessage("表格填充完成，点击生成新 Word 文件保存")

    def extract_seq_text(self, text, fallback):
        detected = detect_number_format(text)
        if detected:
            num_type, sep, has_space, num_part, rest = detected
            return num_part
        return str(fallback)

    # ================== Excel 导入导出 ==================
    def export_to_excel(self):
        if not self.doc or not self.doc.tables:
            QMessageBox.warning(self, "警告", "文档中没有表格")
            return
        dialog = QDialog(self)
        dialog.setWindowTitle("选择要导出的表格")
        layout = QVBoxLayout(dialog)
        list_widget = QListWidget()
        list_widget.setSelectionMode(QListWidget.MultiSelection)
        for i, table in enumerate(self.doc.tables):
            item = QListWidgetItem(f"{i+1}: {get_table_name(table, self.doc, i)}")
            item.setData(Qt.UserRole, i)
            list_widget.addItem(item)
        layout.addWidget(list_widget)
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)
        if dialog.exec_() != QDialog.Accepted:
            return
        selected = [item.data(Qt.UserRole) for item in list_widget.selectedItems()]
        if not selected:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "保存 Excel", "", "Excel 文件 (*.xlsx)")
        if not save_path:
            return
        wb = Workbook()
        wb.remove(wb.active)
        for idx in selected:
            table = self.doc.tables[idx]
            ws = wb.create_sheet(title=f"表格{idx+1}")
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    ws.cell(row=r+1, column=c+1, value=cell.text)
        wb.save(save_path)
        QMessageBox.information(self, "完成", f"已导出到 {save_path}")

    def import_from_excel(self):
        if not self.doc or not self.doc.tables:
            QMessageBox.warning(self, "警告", "文档中没有表格")
            return
        excel_path, _ = QFileDialog.getOpenFileName(self, "选择 Excel", "", "Excel 文件 (*.xlsx)")
        if not excel_path:
            return
        wb = load_workbook(excel_path, data_only=True)
        if not wb.sheetnames:
            QMessageBox.warning(self, "警告", "Excel 为空")
            return
        ws = wb[wb.sheetnames[0]]
        headers = [cell.value for cell in ws[1]]
        data_rows = list(ws.iter_rows(min_row=2, values_only=True))
        if not data_rows:
            QMessageBox.warning(self, "警告", "Excel 没有数据")
            return

        dialog = QDialog(self)
        dialog.setWindowTitle("选择目标表格和列映射")
        layout = QFormLayout(dialog)
        combo_table = QComboBox()
        for i, table in enumerate(self.doc.tables):
            combo_table.addItem(f"{i+1}: {get_table_name(table, self.doc, i)}", i)
        layout.addRow("目标表格：", combo_table)
        combo_word_col = QComboBox()
        table_idx = combo_table.currentData()
        if table_idx is not None:
            table = self.doc.tables[table_idx]
            for c in range(len(table.columns)):
                header = table.rows[0].cells[c].text.strip() if table.rows else f"列{c+1}"
                combo_word_col.addItem(f"{c+1} - {header}", c)
        combo_table.currentIndexChanged.connect(lambda: self.update_import_cols(combo_table, combo_word_col))
        layout.addRow("Word 列：", combo_word_col)
        combo_excel_col = QComboBox()
        for i, h in enumerate(headers):
            combo_excel_col.addItem(f"{i+1} - {h if h else ''}", i)
        layout.addRow("Excel 列：", combo_excel_col)
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addRow(buttons)
        if dialog.exec_() != QDialog.Accepted:
            return
        target_table_idx = combo_table.currentData()
        word_col = combo_word_col.currentData()
        excel_col = combo_excel_col.currentData()
        if target_table_idx is None or word_col is None or excel_col is None:
            return
        table = self.doc.tables[target_table_idx]
        needed = len(data_rows)
        current = len(table.rows) - 1
        if current < needed:
            for _ in range(needed - current):
                table.add_row()
        for i, row_data in enumerate(data_rows):
            value = row_data[excel_col]
            table.rows[i+1].cells[word_col].text = str(value) if value is not None else ""
        self.statusBar().showMessage("导入完成，点击生成新 Word 文件保存")

    def update_import_cols(self, combo_table, combo_word_col):
        combo_word_col.clear()
        table_idx = combo_table.currentData()
        if table_idx is None:
            return
        table = self.doc.tables[table_idx]
        for c in range(len(table.columns)):
            header = table.rows[0].cells[c].text.strip() if table.rows else f"列{c+1}"
            combo_word_col.addItem(f"{c+1} - {header}", c)

# ================== 程序入口 ==================
if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())
